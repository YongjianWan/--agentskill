#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate meeting minutes from transcription. (v2.0)

Usage:
    python generate_minutes.py --input transcription.txt --title "会议主题"

Output:
    meeting_minutes_YYYYMMDD_HHMMSS.docx - Structured meeting minutes
    actions_YYYYMMDD_HHMMSS.json - Action items
"""

import sys
import json
import re
import argparse
from datetime import datetime
from pathlib import Path
from typing import List, Optional
from dataclasses import dataclass, field

# Windows 控制台 UTF-8 编码设置
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    exit(1)


@dataclass
class ActionItem:
    action: str
    owner: str = "待定"
    deadline: Optional[str] = None
    status: str = "待处理"
    source: str = ""
    
    def to_dict(self) -> dict:
        return {
            "action": self.action,
            "owner": self.owner,
            "deadline": self.deadline,
            "status": self.status,
            "source": self.source
        }


@dataclass
class MeetingTopic:
    title: str
    discussion_points: List[str] = field(default_factory=list)
    conclusion: str = ""
    action_items: List[ActionItem] = field(default_factory=list)


@dataclass
class MeetingMinutes:
    title: str = "未命名会议"
    date: str = ""
    time_range: str = ""
    location: str = ""
    participants: List[str] = field(default_factory=list)
    recorder: str = ""
    topics: List[MeetingTopic] = field(default_factory=list)
    risks: List[str] = field(default_factory=list)
    pending_confirmations: List[str] = field(default_factory=list)
    
    def to_dict(self) -> dict:
        return {
            "title": self.title,
            "date": self.date,
            "time_range": self.time_range,
            "location": self.location,
            "participants": self.participants,
            "recorder": self.recorder,
            "topics": [
                {
                    "title": t.title,
                    "discussion_points": t.discussion_points,
                    "conclusion": t.conclusion,
                    "action_items": [a.to_dict() for a in t.action_items]
                }
                for t in self.topics
            ],
            "risks": self.risks,
            "pending_confirmations": self.pending_confirmations
        }


def parse_transcription(text: str) -> tuple:
    """
    Parse transcription with timestamps.
    Returns: (participants_set, segments_list)
    """
    lines = text.strip().split('\n')
    segments = []
    participants = set()
    
    # Pattern: [HH:MM:SS.ms] 发言人：内容
    pattern = r'\[(\d{2}:\d{2}:\d{2}\.\d{3})\]\s*([^：:]+)[：:](.+)'
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        match = re.match(pattern, line)
        if match:
            timestamp = match.group(1)
            speaker = match.group(2).strip()
            content = match.group(3).strip()
            
            segments.append({
                'timestamp': timestamp,
                'speaker': speaker,
                'content': content
            })
            participants.add(speaker)
    
    return participants, segments


def extract_action_items(segments: list, topic_title: str = "") -> List[ActionItem]:
    """Extract action items from segments - improved filtering."""
    actions = []
    
    # Strong action indicators (high confidence)
    strong_action_patterns = [
        r'(负责).+?(完成|提交|准备|制定)',
        r'(完成|提交|准备|制定).+?(?:日前|前|截止)',
        r'(确认|跟进|处理|解决).+?(?:问题|事项|风险)',
    ]
    
    # Deadline patterns
    deadline_patterns = [
        r'(\d{1,2})月(\d{1,2})日(?:前)?',
        r'(下周[一二三四五六日]|本月底|月底|下周)',
    ]
    
    # Non-action filters (exclude these)
    non_action_patterns = [
        r'^(大家|各位|我们|今天)',
        r'(早上好|下午好|大家好)',
        r'(召开|召开.*会议|召开.*会)',
        r'(参会人员?有|人员有)',
        r'^(我在|我是|.*?在)',
        r'^[好的|没问题|确认](?:，)?$',
        r'^那',
    ]
    
    for seg in segments:
        text = seg['content']
        speaker = seg['speaker']
        
        # Skip non-action statements
        is_non_action = any(re.search(p, text) for p in non_action_patterns)
        if is_non_action:
            continue
        
        # Check for strong action patterns
        is_action = any(re.search(p, text) for p in strong_action_patterns)
        has_deadline = any(re.search(p, text) for p in deadline_patterns)
        
        # Must have deadline OR strong action pattern
        if not (is_action or has_deadline):
            continue
        
        # Skip if just self-introduction or confirmation
        if len(text) < 10 or re.match(r'^(确认|好的|没问题)[，。]?$', text):
            continue
        
        # Try to extract owner
        owner = speaker  # Default to current speaker
        
        # Check if text mentions another person as responsible
        person_pattern = r'([\u4e00-\u9fa5]{2,4})(?:负责|完成|提交)'
        person_match = re.search(person_pattern, text)
        if person_match:
            mentioned = person_match.group(1)
            if mentioned in ['张三', '李四', '王五', '主持人', '经理', '主管', '负责人']:
                owner = mentioned
        
        # Extract deadline
        deadline = None
        for dp in deadline_patterns:
            dm = re.search(dp, text)
            if dm:
                if '月' in dm.group(0):
                    deadline = dm.group(0).replace('前', '')
                else:
                    deadline = dm.group(0)
                break
        
        # Clean up action text
        action_text = text.strip()
        action_text = re.sub(r'^[，。、\s]+', '', action_text)  # Remove leading punctuation
        action_text = re.sub(r'[\n\r]+', ' ', action_text)  # Remove newlines
        
        actions.append(ActionItem(
            action=action_text,
            owner=owner,
            deadline=deadline,
            source=topic_title
        ))
    
    return actions


def detect_risks(segments: list) -> List[str]:
    """Detect risk points from segments."""
    risks = []
    risk_keywords = ['风险', '问题', '不稳定', '延迟', '延期', ' blocker', '阻塞', '困难', '挑战']
    
    for seg in segments:
        text = seg['content']
        if any(kw in text for kw in risk_keywords):
            risks.append(f"[{seg['speaker']}] {text}")
    
    return risks


def generate_minutes(segments: list, participants: set, title: str = "") -> MeetingMinutes:
    """Generate structured meeting minutes from parsed segments."""
    minutes = MeetingMinutes(
        title=title or "未命名会议",
        participants=list(participants),
        date=datetime.now().strftime('%Y年%m月%d日')
    )
    
    if not segments:
        return minutes
    
    # Get time range from first and last segment
    first_time = segments[0]['timestamp'][:5]  # HH:MM
    last_time = segments[-1]['timestamp'][:5]
    minutes.time_range = f"{first_time}-{last_time}"
    
    # Group segments by topic (simple approach: by time chunks)
    # Split into topics every ~10 segments or when keyword detected
    topic_segments = []
    current_topic_segs = []
    
    topic_keywords = ['开始', '首先', '第一', '第二', '第三', '接下来', '然后', '关于', '讨论']
    
    for i, seg in enumerate(segments):
        # Check if this is a new topic starter
        is_new_topic = any(kw in seg['content'][:10] for kw in topic_keywords)
        
        if is_new_topic and current_topic_segs and len(current_topic_segs) > 3:
            topic_segments.append(current_topic_segs)
            current_topic_segs = [seg]
        else:
            current_topic_segs.append(seg)
    
    if current_topic_segs:
        topic_segments.append(current_topic_segs)
    
    # If no topics detected, put all in one topic
    if not topic_segments:
        topic_segments = [segments]
    
    # Create topics
    for i, topic_segs in enumerate(topic_segments, 1):
        # Topic title from first segment or keyword
        first_content = topic_segs[0]['content']
        if len(first_content) > 20:
            topic_title = first_content[:20] + "..."
        else:
            topic_title = first_content
        
        # Collect discussion points
        discussion_points = []
        for seg in topic_segs:
            point = f"[{seg['speaker']}] {seg['content']}"
            discussion_points.append(point)
        
        # Extract action items for this topic
        actions = extract_action_items(topic_segs, topic_title)
        
        topic = MeetingTopic(
            title=f"议题{i}：{topic_title}",
            discussion_points=discussion_points,
            action_items=actions
        )
        minutes.topics.append(topic)
    
    # Detect risks
    minutes.risks = detect_risks(segments)
    
    return minutes


def create_docx(minutes: MeetingMinutes, output_path: str):
    """Create Word document for meeting minutes."""
    doc = Document()
    
    # Title
    title = doc.add_heading('会议纪要', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Basic info
    doc.add_heading('一、会议基本信息', level=1)
    info = doc.add_paragraph()
    info.add_run(f"会议主题：{minutes.title}\n")
    info.add_run(f"会议时间：{minutes.date} {minutes.time_range}\n")
    info.add_run(f"会议地点：{minutes.location or '待定'}\n")
    info.add_run(f"参会人员：{'、'.join(minutes.participants) if minutes.participants else '待定'}\n")
    info.add_run(f"记录人员：{minutes.recorder or '待定'}\n")
    
    # Topics
    for topic in minutes.topics:
        doc.add_heading(topic.title, level=1)
        
        # Discussion points
        if topic.discussion_points:
            doc.add_heading('讨论要点', level=2)
            for point in topic.discussion_points[:10]:  # Limit points
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(point[:300])  # Limit length
        
        # Action items
        if topic.action_items:
            doc.add_heading('行动项', level=2)
            action_table = doc.add_table(rows=1, cols=4)
            action_table.style = 'Light List Accent 1'
            
            # Header
            hdr_cells = action_table.rows[0].cells
            hdr_cells[0].text = '序号'
            hdr_cells[1].text = '行动事项'
            hdr_cells[2].text = '负责人'
            hdr_cells[3].text = '完成期限'
            
            # Data rows
            for j, action in enumerate(topic.action_items, 1):
                row_cells = action_table.add_row().cells
                row_cells[0].text = str(j)
                row_cells[1].text = action.action
                row_cells[2].text = action.owner
                row_cells[3].text = action.deadline or '待定'
        
        doc.add_paragraph()
    
    # Risks
    if minutes.risks:
        doc.add_heading('二、风险点', level=1)
        for risk in minutes.risks:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(risk)
    
    # Save
    doc.save(output_path)
    print(f"Created: {output_path}")


def create_actions_json(minutes: MeetingMinutes, output_path: str):
    """Create actions.json file."""
    all_actions = []
    for topic in minutes.topics:
        for action in topic.action_items:
            action_dict = action.to_dict()
            action_dict['topic'] = topic.title
            all_actions.append(action_dict)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(all_actions, f, ensure_ascii=False, indent=2)
    
    print(f"Created: {output_path} ({len(all_actions)} action items)")


def main():
    parser = argparse.ArgumentParser(description="Generate meeting minutes v2.0")
    parser.add_argument("--input", "-i", required=True, help="Input transcription file (with timestamps)")
    parser.add_argument("--output-dir", "-o", default="../output", help="Output directory (default: ../output)")
    parser.add_argument("--title", "-t", default="", help="Meeting title")
    args = parser.parse_args()
    
    # Read input
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: File not found: {input_path}")
        return
    
    with open(input_path, 'r', encoding='utf-8') as f:
        text = f.read()
    
    # Parse transcription
    print("Parsing transcription...")
    participants, segments = parse_transcription(text)
    print(f"Found {len(segments)} segments, {len(participants)} participants")
    
    if not segments:
        print("Error: No valid segments found in input file")
        return
    
    # Generate minutes
    print("Generating meeting minutes...")
    minutes = generate_minutes(segments, participants, args.title)
    
    # Create output directory (resolve relative to script location)
    script_dir = Path(__file__).parent
    output_dir = script_dir / args.output_dir
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Generate timestamp
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    
    # Create outputs
    docx_path = output_dir / f"meeting_minutes_{timestamp}.docx"
    create_docx(minutes, str(docx_path))
    
    json_path = output_dir / f"actions_{timestamp}.json"
    create_actions_json(minutes, str(json_path))
    
    # Also save structured JSON
    raw_json_path = output_dir / f"meeting_minutes_{timestamp}.json"
    with open(raw_json_path, 'w', encoding='utf-8') as f:
        json.dump(minutes.to_dict(), f, ensure_ascii=False, indent=2)
    print(f"Created: {raw_json_path}")
    
    # Summary
    total_actions = sum(len(t.action_items) for t in minutes.topics)
    print(f"\n[OK] Generated {len(minutes.topics)} topics, {total_actions} action items")
    if minutes.risks:
        print(f"     Found {len(minutes.risks)} risk points")


if __name__ == "__main__":
    main()

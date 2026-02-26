#!/usr/bin/env python3
"""
Meeting Management Skill - AI Callable Interface

设计原则：
- Skill 层（本模块）：负责数据层（转写、存储、I/O、格式转换）
- AI 层（调用方）：负责智能理解（议题提取、结论识别、行动项抽取）

主要接口：
    transcribe()              - 音频转写（Skill）
    create_meeting_skeleton() - 创建会议骨架（Skill，AI需填充内容）
    save_meeting()            - 保存会议（Skill）
    query_meetings()          - 查询历史（Skill）
    update_meeting()          - 更新会议（Skill）

Usage:
    from meeting_skill import transcribe, create_meeting_skeleton, Meeting, Topic, ActionItem, save_meeting
    
    # Step 1: 转写（Skill）
    result = transcribe("meeting.mp3")
    
    # Step 2: 创建骨架 & AI 填充（AI 层）
    meeting = create_meeting_skeleton(result["full_text"])
    meeting.title = "AI 识别的标题"
    meeting.topics = [Topic(title="议题1", discussion_points=[...], conclusion="...", action_items=[...])]
    
    # Step 3: 保存（Skill）
    files = save_meeting(meeting, output_dir="./output")

Note:
    generate_minutes() 已弃用，请使用 create_meeting_skeleton()
"""

import json
import re
import warnings
import hashlib
from pathlib import Path
from datetime import datetime
from typing import List, Optional, Dict, Any, Tuple
from dataclasses import dataclass, field

warnings.filterwarnings("ignore")


# ============ 数据模型 ============

@dataclass
class TranscriptionSegment:
    timestamp: str
    speaker: str
    text: str


@dataclass
class PolicyRef:
    """政策引用（设计预留，未来功能）"""
    policy_id: str = ""           # 政策ID
    clause_ref: str = ""          # 引用条款
    check_required: bool = False  # 是否待核查
    
    def to_dict(self) -> Dict:
        return {
            "policy_id": self.policy_id,
            "clause_ref": self.clause_ref,
            "check_required": self.check_required
        }


@dataclass
class EnterpriseRef:
    """企业关联（设计预留，未来功能）"""
    enterprise_id: str = ""       # 企业ID
    cooperation_item: str = ""    # 合作事项
    contact_person: str = ""      # 对接人
    contact_permission: str = ""  # 权限级别
    
    def to_dict(self) -> Dict:
        return {
            "enterprise_id": self.enterprise_id,
            "cooperation_item": self.cooperation_item,
            "contact_person": self.contact_person,
            "contact_permission": self.contact_permission
        }


@dataclass
class ProjectRef:
    """项目关联（设计预留，未来功能）"""
    project_id: str = ""          # 项目ID
    milestone: str = ""           # 涉及里程碑
    change_point: str = ""        # 变更点说明
    
    def to_dict(self) -> Dict:
        return {
            "project_id": self.project_id,
            "milestone": self.milestone,
            "change_point": self.change_point
        }


@dataclass
class ActionItem:
    action: str
    owner: str
    deadline: str
    deliverable: str = ""
    status: str = "待处理"
    source_topic: str = ""
    # 关联实体（AI 识别填充）
    related_policy: Optional[PolicyRef] = None      # 关联政策
    related_enterprise: Optional[EnterpriseRef] = None  # 关联企业
    related_project: Optional[ProjectRef] = None    # 关联项目
    
    def to_dict(self) -> Dict:
        return {
            "action": self.action,
            "owner": self.owner,
            "deadline": self.deadline,
            "deliverable": self.deliverable,
            "status": self.status,
            "source_topic": self.source_topic,
            "related_policy": self.related_policy.to_dict() if self.related_policy else None,
            "related_enterprise": self.related_enterprise.to_dict() if self.related_enterprise else None,
            "related_project": self.related_project.to_dict() if self.related_project else None
        }


@dataclass
class Topic:
    title: str
    discussion_points: List[str]
    conclusion: str = ""  # 新增：结论
    uncertain: List[str] = field(default_factory=list)  # 新增：不确定内容
    action_items: List[ActionItem] = field(default_factory=list)
    
    def to_dict(self) -> Dict:
        return {
            "title": self.title,
            "discussion_points": self.discussion_points,
            "conclusion": self.conclusion,
            "uncertain": self.uncertain,
            "action_items": [a.to_dict() for a in self.action_items]
        }


@dataclass
class Meeting:
    id: str
    title: str
    date: str
    time_range: str
    location: str
    participants: List[str]
    recorder: str
    topics: List[Topic]
    risks: List[str]
    pending_confirmations: List[str]  # 新增：待确认事项
    audio_path: Optional[str] = None
    version: int = 1  # 新增：版本号
    created_at: str = field(default_factory=lambda: datetime.now().isoformat())
    updated_at: str = field(default_factory=lambda: datetime.now().isoformat())
    
    # === 关联实体（设计预留，AI 识别填充）===
    # 会议涉及的政策引用
    policy_refs: List[PolicyRef] = field(default_factory=list)
    # 会议涉及的企业
    enterprise_refs: List[EnterpriseRef] = field(default_factory=list)
    # 会议关联的项目
    project_refs: List[ProjectRef] = field(default_factory=list)
    
    # === 状态管理 ===
    status: str = "draft"  # draft/final/archived
    
    def to_dict(self) -> Dict:
        return {
            "id": self.id,
            "title": self.title,
            "date": self.date,
            "time_range": self.time_range,
            "location": self.location,
            "participants": self.participants,
            "recorder": self.recorder,
            "topics": [t.to_dict() for t in self.topics],
            "risks": self.risks,
            "pending_confirmations": self.pending_confirmations,
            "audio_path": self.audio_path,
            "version": self.version,
            "created_at": self.created_at,
            "updated_at": self.updated_at,
            # 关联实体
            "policy_refs": [p.to_dict() for p in self.policy_refs],
            "enterprise_refs": [e.to_dict() for e in self.enterprise_refs],
            "project_refs": [p.to_dict() for p in self.project_refs],
            "status": self.status
        }


# ============ 核心功能 ============

def transcribe(audio_path: str, model: str = "auto", language: str = "zh") -> Dict[str, Any]:
    """
    转写音频文件为结构化文本
    
    Args:
        audio_path: 音频文件路径
        model: 模型大小 (tiny/base/small/medium/large/auto)。默认 auto 自动选择：
               - 会议音频自动使用 small 模型（推荐，效果更好）
        language: 语言代码
    
    Returns:
        {
            "segments": [{"timestamp": "00:00:01", "speaker": "...", "text": "..."}],
            "full_text": "...",
            "participants": ["..."],
            "duration": 1800,
            "model_used": "small"  # 实际使用的模型
        }
    """
    from faster_whisper import WhisperModel
    
    # 自动选择模型：会议场景默认用 small（效果更好）
    if model == "auto":
        # 会议场景推荐 small（244MB，中文识别效果更好）
        model = "small"
        print(f"[Info] 会议音频自动使用 small 模型（识别效果更好）")
    
    if not Path(audio_path).exists():
        raise FileNotFoundError(f"Audio file not found: {audio_path}")
    
    model_obj = WhisperModel(model, device="cpu", compute_type="int8")
    segments, info = model_obj.transcribe(audio_path, beam_size=5, language=language)
    
    result_segments = []
    speakers = set()
    full_text_parts = []
    
    for i, segment in enumerate(segments):
        start_time = segment.start
        text = segment.text.strip()
        
        hours = int(start_time // 3600)
        minutes = int((start_time % 3600) // 60)
        seconds = int(start_time % 60)
        time_str = f"{hours:02d}:{minutes:02d}:{seconds:02d}"
        
        # 发言人识别（简单启发式）
        speaker = f"Speaker{i % 3 + 1}"
        if "：" in text or ":" in text:
            parts = text.split("：", 1) if "：" in text else text.split(":", 1)
            if len(parts) == 2 and len(parts[0]) < 10:
                speaker = parts[0].strip()
                text = parts[1].strip()
        
        speakers.add(speaker)
        
        result_segments.append({
            "timestamp": time_str,
            "speaker": speaker,
            "text": text
        })
        
        full_text_parts.append(f"[{time_str}] {speaker}: {text}")
    
    return {
        "segments": result_segments,
        "full_text": "\n".join(full_text_parts),
        "participants": list(speakers),
        "duration": int(info.duration) if info.duration else 0,
        "language": info.language,
        "model_used": model  # 返回实际使用的模型
    }


def create_meeting_skeleton(
    transcription: str,
    meeting_id: Optional[str] = None,
    title: str = "",
    date: Optional[str] = None,
    time_range: str = "",
    location: str = "",
    participants: Optional[List[str]] = None,
    recorder: str = "",
    audio_path: Optional[str] = None,
    version: int = 1
) -> Meeting:
    """
    创建会议骨架结构（基础数据层，不做智能理解）
    
    本函数只负责：
    - 解析转写文本（时间戳/发言人/内容分段）
    - 基础发言人大致识别（需 AI 校验）
    - 创建空的 Meeting 结构供 AI 填充
    
    AI 需要填充：
    - title: 会议标题（语义识别）
    - topics: 议题列表（议题提取、结论识别、行动项抽取）
    - risks: 风险点（语义识别）
    - pending_confirmations: 待确认事项
    
    Args:
        transcription: 转写文本（transcribe() 的 full_text 输出）
        meeting_id: 会议ID（可选，自动生成）
        title: 会议标题（可选，AI 应重新识别）
        date: 日期（可选，默认今天）
        time_range: 时间范围
        location: 地点
        participants: 参会人列表（可选，会从文本基础提取）
        recorder: 记录人
        audio_path: 音频文件路径
        version: 版本号
    
    Returns:
        Meeting 对象（topics 为空列表，需 AI 填充）
    
    Example:
        >>> result = transcribe("meeting.mp3")
        >>> meeting = create_meeting_skeleton(result["full_text"])
        >>> # AI 填充内容
        >>> meeting.title = "产品评审会"
        >>> meeting.topics = [Topic(title="议题1", ...)]
        >>> save_meeting(meeting)
    """
    if date is None:
        date = datetime.now().strftime("%Y-%m-%d")
    
    # 生成会议ID
    if meeting_id is None:
        meeting_id = _generate_meeting_id(date, title or "未命名会议")
    
    # 解析转写文本（基础清洗）
    segments = _parse_transcription(transcription)
    
    # 基础参会人提取（简单启发式，AI 应校验修正）
    detected_participants = participants or _extract_participants_basic(segments)
    
    # === 注意：不做智能理解，topics 留空供 AI 填充 ===
    # 如果用户提供了 title，保留；否则设为占位符，AI 应重新识别
    if not title:
        title = "【AI 待识别标题】"
    
    # 创建空的 topics，AI 应填充
    # 可选：将原始分段作为参考信息（非结构化）
    _ = [seg["text"] for seg in segments[:5]] if segments else ["无内容"]  # 保留用于调试
    
    return Meeting(
        id=meeting_id,
        title=title,
        date=date,
        time_range=time_range,
        location=location,
        participants=detected_participants,
        recorder=recorder,
        topics=[],  # ← AI 负责填充议题
        risks=[],   # ← AI 负责识别风险
        pending_confirmations=[],  # ← AI 负责标记待确认
        audio_path=audio_path,
        version=version
    )


def generate_minutes(
    transcription: str,
    meeting_id: Optional[str] = None,
    title: str = "",
    date: Optional[str] = None,
    time_range: str = "",
    location: str = "",
    participants: Optional[List[str]] = None,
    recorder: str = "",
    audio_path: Optional[str] = None,
    version: int = 1,
    use_ai: bool = True
) -> Meeting:
    """
    生成会议纪要（现在使用 DeepSeek AI）
    
    Args:
        transcription: 转写文本
        title: 会议标题
        use_ai: 是否使用 AI 生成（默认 True）
        ... 其他参数
        
    Returns:
        Meeting 对象（topics 已由 AI 填充）
    """
    if date is None:
        date = datetime.now().strftime("%Y-%m-%d")
    
    # 生成会议ID
    if meeting_id is None:
        meeting_id = _generate_meeting_id(date, title or "未命名会议")
    
    # 基础参会人提取
    segments = _parse_transcription(transcription)
    detected_participants = participants or _extract_participants_basic(segments)
    
    # 使用 AI 生成纪要内容
    if use_ai:
        try:
            from .ai_minutes_generator import generate_minutes_with_ai
            
            ai_result = generate_minutes_with_ai(transcription, title_hint=title)
            
            if ai_result:
                # 使用 AI 生成的标题（如果未提供）
                ai_title = ai_result.get("title", title or "未命名会议")
                
                # 转换 topics
                ai_topics = []
                for t in ai_result.get("topics", []):
                    action_items = []
                    for a in t.get("action_items", []):
                        action_items.append(ActionItem(
                            action=a.get("action", ""),
                            owner=a.get("owner", "待定"),
                            deadline=a.get("deadline", ""),
                            deliverable="",
                            status="待处理"
                        ))
                    
                    ai_topics.append(Topic(
                        title=t.get("title", "未命名议题"),
                        discussion_points=t.get("discussion_points", []),
                        conclusion=t.get("conclusion", ""),
                        uncertain=[],
                        action_items=action_items
                    ))
                
                # 使用 AI 识别的参会人（如果有）
                ai_participants = ai_result.get("participants", detected_participants)
                
                meeting = Meeting(
                    id=meeting_id,
                    title=ai_title,
                    date=date,
                    time_range=time_range,
                    location=location,
                    participants=ai_participants,
                    recorder=recorder,
                    topics=ai_topics,
                    risks=ai_result.get("risks", []),
                    pending_confirmations=ai_result.get("pending_confirmations", []),
                    audio_path=audio_path,
                    version=version
                )
                
                # 标记为 AI 生成
                meeting.status = "ai_generated"
                return meeting
                
        except Exception as e:
            print(f"[generate_minutes] AI 生成失败，降级到骨架模式: {e}")
    
    # 降级：返回骨架（AI 失败或 use_ai=False）
    return create_meeting_skeleton(
        transcription=transcription,
        meeting_id=meeting_id,
        title=title,
        date=date,
        time_range=time_range,
        location=location,
        participants=detected_participants,
        recorder=recorder,
        audio_path=audio_path,
        version=version
    )


def save_meeting(
    meeting: Meeting,
    output_dir: str = "./output",
    create_version: bool = True
) -> Dict[str, str]:
    """
    保存会议记录（支持版本控制）
    
    Args:
        meeting: Meeting 对象
        output_dir: 输出目录
        create_version: 是否创建版本目录
    
    Returns:
        {"docx": "...", "json": "...", "audio_backup": "..."}
    """
    from docx import Document  # noqa: F401
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: F401
    
    # 构建输出路径
    if create_version:
        meeting_dir = Path(output_dir) / "meetings" / meeting.date[:4] / meeting.date[5:7] / meeting.id
    else:
        meeting_dir = Path(output_dir)
    
    meeting_dir.mkdir(parents=True, exist_ok=True)
    
    # 保存 JSON（含版本号）
    json_path = meeting_dir / f"minutes_v{meeting.version}.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(meeting.to_dict(), f, ensure_ascii=False, indent=2)
    
    # 保存 Word
    docx_path = meeting_dir / f"minutes_v{meeting.version}.docx"
    _create_word_document(meeting, docx_path)
    
    # 备份录音
    audio_backup_path = None
    if meeting.audio_path and Path(meeting.audio_path).exists():
        audio_name = Path(meeting.audio_path).name
        audio_backup_path = meeting_dir / f"audio_{audio_name}"
        import shutil
        shutil.copy2(meeting.audio_path, audio_backup_path)
    
    # 更新最新版本链接
    latest_json = meeting_dir / "minutes_latest.json"
    latest_docx = meeting_dir / "minutes_latest.docx"
    
    # Windows 不支持 symlink，直接复制
    import shutil
    shutil.copy2(json_path, latest_json)
    shutil.copy2(docx_path, latest_docx)
    
    # 保存到全局台账
    _append_to_action_registry(meeting, output_dir)
    
    result = {
        "json": str(json_path),
        "docx": str(docx_path),
        "meeting_dir": str(meeting_dir)
    }
    
    if audio_backup_path:
        result["audio_backup"] = str(audio_backup_path)
    
    return result


def update_meeting(
    meeting_id: str,
    output_dir: str = "./output",
    **updates
) -> Meeting:
    """
    更新会议纪要（创建新版本）
    
    Args:
        meeting_id: 会议ID
        output_dir: 输出目录
        **updates: 要更新的字段
    
    Returns:
        新版本 Meeting 对象
    """
    # 查找最新版本
    meeting_dir = _find_meeting_dir(meeting_id, output_dir)
    if not meeting_dir:
        raise FileNotFoundError(f"Meeting {meeting_id} not found")
    
    latest_json = meeting_dir / "minutes_latest.json"
    
    # 加载现有数据
    with open(latest_json, "r", encoding="utf-8") as f:
        data = json.load(f)
    
    # 创建新版本
    new_version = data.get("version", 1) + 1
    data["version"] = new_version
    data["updated_at"] = datetime.now().isoformat()
    
    # 应用更新
    for key, value in updates.items():
        if key in data:
            data[key] = value
    
    # 重建 Meeting 对象
    def _rebuild_topic(t_data: dict) -> Topic:
        """Rebuild Topic from dict, handling nested ActionItem"""
        # Handle action_items
        action_items_data = t_data.get("action_items", [])
        action_items = []
        for a_data in action_items_data:
            if isinstance(a_data, dict):
                # Filter valid fields for ActionItem
                valid_fields = {"action", "owner", "deadline", "deliverable", "status", "source_topic", 
                               "related_policy", "related_enterprise", "related_project"}
                filtered = {k: v for k, v in a_data.items() if k in valid_fields}
                action_items.append(ActionItem(**filtered))
            elif isinstance(a_data, ActionItem):
                action_items.append(a_data)
        
        return Topic(
            title=t_data.get("title", ""),
            discussion_points=t_data.get("discussion_points", []),
            conclusion=t_data.get("conclusion", ""),
            uncertain=t_data.get("uncertain", []),
            action_items=action_items
        )
    
    topics = [_rebuild_topic(t) for t in data.get("topics", [])]
    meeting = Meeting(
        id=data["id"],
        title=data["title"],
        date=data["date"],
        time_range=data.get("time_range", ""),
        location=data.get("location", ""),
        participants=data.get("participants", []),
        recorder=data.get("recorder", ""),
        topics=topics,
        risks=data.get("risks", []),
        pending_confirmations=data.get("pending_confirmations", []),
        audio_path=data.get("audio_path"),
        version=new_version,
        created_at=data.get("created_at", datetime.now().isoformat()),
        updated_at=data["updated_at"]
    )
    
    # 保存新版本
    save_meeting(meeting, output_dir, create_version=True)
    
    return meeting


def query_meetings(
    output_dir: str = "./output",
    date_range: Optional[Tuple[str, str]] = None,
    participants: Optional[List[str]] = None,
    keywords: Optional[List[str]] = None
) -> List[Dict]:
    """
    查询历史会议
    
    Args:
        output_dir: 输出目录
        date_range: (开始日期, 结束日期) 格式 YYYY-MM-DD
        participants: 参会人列表
        keywords: 关键词列表
    
    Returns:
        Meeting 列表
    """
    meetings_dir = Path(output_dir) / "meetings"
    results = []
    
    if not meetings_dir.exists():
        return results
    
    # 遍历所有会议目录
    for year_dir in meetings_dir.iterdir():
        if not year_dir.is_dir():
            continue
        for month_dir in year_dir.iterdir():
            if not month_dir.is_dir():
                continue
            for meeting_dir in month_dir.iterdir():
                if not meeting_dir.is_dir():
                    continue
                
                latest_json = meeting_dir / "minutes_latest.json"
                if not latest_json.exists():
                    continue
                
                try:
                    with open(latest_json, "r", encoding="utf-8") as f:
                        data = json.load(f)
                    
                    # 过滤条件
                    if date_range:
                        if not (date_range[0] <= data.get("date", "") <= date_range[1]):
                            continue
                    
                    if participants:
                        if not any(p in data.get("participants", []) for p in participants):
                            continue
                    
                    if keywords:
                        text = json.dumps(data, ensure_ascii=False)
                        if not any(kw in text for kw in keywords):
                            continue
                    
                    results.append(data)
                except Exception:
                    continue
    
    # 按日期排序
    results.sort(key=lambda x: x.get("date", ""), reverse=True)
    return results


# ============ 内部辅助函数 ============

def _generate_meeting_id(date: str, title: str) -> str:
    """生成会议ID"""
    timestamp = datetime.now().strftime("%H%M%S")
    hash_input = f"{date}_{title}_{timestamp}"
    hash_suffix = hashlib.md5(hash_input.encode()).hexdigest()[:6]
    return f"M{date.replace('-', '')}_{timestamp}_{hash_suffix}"


def _parse_transcription(text: str) -> List[Dict[str, str]]:
    """解析转写文本为段落列表"""
    segments = []
    lines = text.strip().split("\n")
    pattern = r"\[(\d{2}:\d{2}:\d{2})\]\s*([^：:]+)[：:](.+)"
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        match = re.match(pattern, line)
        if match:
            segments.append({
                "time": match.group(1),
                "speaker": match.group(2).strip(),
                "text": match.group(3).strip()
            })
        else:
            segments.append({"time": "", "speaker": "", "text": line})
    
    return segments


def _extract_participants_basic(segments: List[Dict]) -> List[str]:
    """
    基础参会人提取（简单启发式）
    
    仅根据转写文本中的发言人标记提取，可能存在：
    - 遗漏（未发言的人）
    - 噪音（错误识别的人名）
    - 重复（同一人多种称呼）
    
    AI 应结合上下文进行校验和修正。
    """
    speakers = set()
    for seg in segments:
        if seg.get("speaker"):
            # 清理可能的噪音
            speaker = seg["speaker"].strip()
            if speaker and len(speaker) < 20:  # 简单长度过滤
                speakers.add(speaker)
    return list(speakers) or ["【AI 待识别参会人】"]



    
    # 默认取前20字
    return first_text[:20] if len(first_text) <= 20 else first_text[:20] + "..."


def _create_word_document(meeting: Meeting, docx_path: Path):
    """创建 Word 文档"""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    
    doc = Document()
    
    # 标题
    title = doc.add_heading(meeting.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 版本信息
    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    version_run = version_para.add_run(f"版本: v{meeting.version}")
    version_run.font.size = Pt(9)
    version_run.font.color.rgb = None  # 灰色
    
    # 基本信息
    doc.add_heading("一、会议基本信息", level=1)
    info = doc.add_paragraph()
    info.add_run(f"会议主题：{meeting.title}\n")
    info.add_run(f"会议时间：{meeting.date} {meeting.time_range}\n")
    info.add_run(f"会议地点：{meeting.location or '待定'}\n")
    info.add_run(f"参会人员：{'、'.join(meeting.participants)}\n")
    info.add_run(f"记录人员：{meeting.recorder or '待定'}\n")
    
    # 议题
    doc.add_heading("二、议题与讨论", level=1)
    
    for i, topic in enumerate(meeting.topics, 1):
        doc.add_heading(f"议题{i}：{topic.title}", level=2)
        
        # 讨论要点
        if topic.discussion_points:
            doc.add_heading("（一）讨论要点", level=3)
            for point in topic.discussion_points[:5]:
                doc.add_paragraph(point, style="List Bullet")
        
        # 结论
        if topic.conclusion:
            doc.add_heading("（二）结论", level=3)
            doc.add_paragraph(topic.conclusion)
        
        # 不确定内容
        if topic.uncertain:
            doc.add_heading("（三）待确认事项", level=3)
            for u in topic.uncertain:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(u).font.color.rgb = None  # 标红或标记
        
        # 行动项
        if topic.action_items:
            doc.add_heading("（四）行动项", level=3)
            table = doc.add_table(rows=1, cols=5)
            table.style = "Light List Accent 1"
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "行动事项"
            hdr_cells[1].text = "负责人"
            hdr_cells[2].text = "完成期限"
            hdr_cells[3].text = "交付物"
            hdr_cells[4].text = "状态"
            
            for action in topic.action_items:
                row_cells = table.add_row().cells
                row_cells[0].text = action.action[:50]
                row_cells[1].text = action.owner
                row_cells[2].text = action.deadline
                row_cells[3].text = action.deliverable or "-"
                row_cells[4].text = action.status
    
    # 全局待确认事项
    if meeting.pending_confirmations:
        doc.add_heading("三、待确认事项汇总", level=1)
        for item in meeting.pending_confirmations:
            doc.add_paragraph(item, style="List Bullet")
    
    # 风险点
    if meeting.risks:
        doc.add_heading("四、风险点", level=1)
        for risk in meeting.risks:
            doc.add_paragraph(risk, style="List Bullet")
    
    # 附件
    if meeting.audio_path:
        doc.add_heading("五、附件", level=1)
        doc.add_paragraph(f"录音文件：{Path(meeting.audio_path).name}")
    
    doc.save(str(docx_path))


def _find_meeting_dir(meeting_id: str, output_dir: str) -> Optional[Path]:
    """查找会议目录"""
    meetings_dir = Path(output_dir) / "meetings"
    
    if not meetings_dir.exists():
        return None
    
    # 遍历查找
    for year_dir in meetings_dir.iterdir():
        for month_dir in year_dir.iterdir():
            for meeting_dir in month_dir.iterdir():
                if meeting_dir.name == meeting_id:
                    return meeting_dir
    
    return None


def _append_to_action_registry(meeting: Meeting, output_dir: str):
    """追加行动项到全局台账"""
    registry_path = Path(output_dir) / "action_registry.json"
    
    # 读取现有台账
    registry = []
    if registry_path.exists():
        try:
            with open(registry_path, "r", encoding="utf-8") as f:
                registry = json.load(f)
        except:
            registry = []
    
    # 添加本次行动项
    for topic in meeting.topics:
        for action in topic.action_items:
            registry.append({
                "meeting_id": meeting.id,
                "meeting_title": meeting.title,
                "meeting_date": meeting.date,
                "action": action.action,
                "owner": action.owner,
                "deadline": action.deadline,
                "deliverable": action.deliverable,
                "status": action.status,
                "source_topic": topic.title
            })
    
    # 保存
    with open(registry_path, "w", encoding="utf-8") as f:
        json.dump(registry, f, ensure_ascii=False, indent=2)


# ============ 音频流处理（新增） ============

import tempfile
import time
from typing import BinaryIO

# 全局会话管理器：meeting_id -> session数据
_audio_sessions: Dict[str, dict] = {}

# Whisper模型缓存（延迟加载）
_whisper_model = None


def _get_whisper_model():
    """获取Whisper模型（延迟加载）"""
    global _whisper_model
    if _whisper_model is None:
        from faster_whisper import WhisperModel
        _whisper_model = WhisperModel("small", device="cpu", compute_type="int8")
    return _whisper_model


def init_meeting_session(meeting_id: str, title: str = "", user_id: str = "anonymous", db_session=None) -> str:
    """
    初始化会议会话
    
    Args:
        meeting_id: 会议ID
        title: 会议标题
        user_id: 用户ID
        db_session: 数据库会话（SQLAlchemy AsyncSession）
    
    Returns:
        audio文件路径
    """
    from models.meeting import MeetingModel, MeetingStatus
    
    # 创建目录 output/meetings/YYYY/MM/{meeting_id}/
    now = datetime.now()
    meeting_dir = Path("./output/meetings") / now.strftime("%Y") / now.strftime("%m") / meeting_id
    meeting_dir.mkdir(parents=True, exist_ok=True)
    
    audio_path = meeting_dir / "audio.webm"
    
    # 初始化音频文件（空文件，追加模式后续写入）
    audio_path.touch()
    
    # 内存会话记录
    _audio_sessions[meeting_id] = {
        "audio_path": str(audio_path),
        "meeting_dir": str(meeting_dir),
        "title": title,
        "user_id": user_id,
        "start_time": now,
        "last_chunk_time": 0,  # 上次转写时间戳
        "chunk_count": 0,
        "file_handle": None,  # 懒加载
    }
    
    # 数据库插入记录（如果提供了db_session）
    if db_session is not None:
        # 注意：db_session是异步的，需要await，这里只同步创建记录对象
        # 实际插入由调用方完成
        pass
    
    return str(audio_path)


def _get_file_handle(meeting_id: str) -> BinaryIO:
    """获取音频文件句柄（懒加载，追加模式）"""
    session = _audio_sessions.get(meeting_id)
    if session is None:
        raise ValueError(f"Meeting session not found: {meeting_id}")
    
    if session["file_handle"] is None or session["file_handle"].closed:
        session["file_handle"] = open(session["audio_path"], "ab")
    
    return session["file_handle"]


def transcribe_bytes(audio_bytes: bytes, mime_type: str = "audio/webm") -> Dict[str, Any]:
    """
    直接转写音频bytes（不落盘临时文件）
    
    Args:
        audio_bytes: 音频数据
        mime_type: 音频格式
    
    Returns:
        {"segments": [...], "full_text": "...", "language": "zh"}
    """
    model = _get_whisper_model()
    
    # faster-whisper需要文件路径，用临时文件
    suffix = ".webm" if "webm" in mime_type else ".wav"
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as f:
        f.write(audio_bytes)
        temp_path = f.name
    
    try:
        segments, info = model.transcribe(temp_path, beam_size=5, language="zh")
        results = []
        full_text_parts = []
        
        for seg in segments:
            text = seg.text.strip()
            results.append({
                "start": seg.start,
                "end": seg.end,
                "text": text
            })
            full_text_parts.append(text)
        
        return {
            "segments": results,
            "full_text": " ".join(full_text_parts),
            "language": info.language
        }
    finally:
        Path(temp_path).unlink(missing_ok=True)


def append_audio_chunk(meeting_id: str, chunk_bytes: bytes, sequence: int, 
                       db_session=None) -> Optional[str]:
    """
    追加音频块并触发转写（每30秒触发一次）
    
    Args:
        meeting_id: 会议ID
        chunk_bytes: 音频块数据
        sequence: 块序号
        db_session: 数据库会话
    
    Returns:
        转写文本（如触发转写），否则None
    """
    from models.meeting import TranscriptModel
    
    session = _audio_sessions.get(meeting_id)
    if session is None:
        raise ValueError(f"Meeting session not found: {meeting_id}")
    
    # 追加写入音频文件
    f = _get_file_handle(meeting_id)
    f.write(chunk_bytes)
    f.flush()
    
    session["chunk_count"] += 1
    current_time = time.time()
    
    # 检查是否触发转写（每30秒一次）
    time_since_last = current_time - session["last_chunk_time"]
    should_transcribe = (session["last_chunk_time"] == 0) or (time_since_last >= 30)
    
    transcript_text = None
    
    if should_transcribe:
        # 读取音频文件内容
        f.flush()
        # 重新打开读取（避免写入缓冲问题）
        f.close()
        session["file_handle"] = None
        
        with open(session["audio_path"], "rb") as rf:
            audio_data = rf.read()
        
        if len(audio_data) > 0:
            # 转写
            result = transcribe_bytes(audio_data)
            transcript_text = result.get("full_text", "")
            
            # 记录转写时间
            session["last_chunk_time"] = current_time
            
            # 保存到数据库（如果提供了db_session）
            if db_session is not None:
                # 异步操作需要await，这里返回数据给调用方处理
                pass
        
        # 重新打开文件追加
        session["file_handle"] = open(session["audio_path"], "ab")
    
    return transcript_text


def finalize_meeting(meeting_id: str, db_session=None) -> dict:
    """
    结束会议，全量转写剩余内容，生成纪要，导出Word
    
    Args:
        meeting_id: 会议ID
        db_session: 数据库会话
    
    Returns:
        {"audio_path": "...", "minutes_path": "...", "full_text": "..."}
    """
    from models.meeting import MeetingModel, MeetingStatus
    
    session = _audio_sessions.get(meeting_id)
    if session is None:
        raise ValueError(f"Meeting session not found: {meeting_id}")
    
    # 关闭文件句柄
    if session["file_handle"] and not session["file_handle"].closed:
        session["file_handle"].close()
    session["file_handle"] = None
    
    # 全量转写
    audio_path = Path(session["audio_path"])
    with open(audio_path, "rb") as f:
        audio_data = f.read()
    
    full_transcript = ""
    if len(audio_data) > 0:
        result = transcribe_bytes(audio_data)
        full_transcript = result.get("full_text", "")
    
    # 生成会议纪要（调用AI）
    minutes_path = Path(session["meeting_dir"]) / "minutes.docx"
    
    # 使用generate_minutes生成纪要（需要转写文本）
    meeting_data = generate_minutes(
        transcription=full_transcript,
        meeting_id=meeting_id,
        title=session["title"],
        date=session["start_time"].strftime("%Y-%m-%d"),
        audio_path=str(audio_path)
    )
    
    # 保存Word
    files = save_meeting(meeting_data, output_dir="./output", create_version=False)
    
    # 移动/重命名为标准路径
    if "docx" in files:
        import shutil
        src_docx = files["docx"]
        if Path(src_docx) != minutes_path:
            shutil.move(src_docx, minutes_path)
    
    # 清理会话
    del _audio_sessions[meeting_id]
    
    return {
        "meeting_id": meeting_id,
        "audio_path": str(audio_path),
        "minutes_path": str(minutes_path),
        "full_text": full_transcript,
        "chunk_count": session["chunk_count"]
    }


# ============ CLI ============

if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("Usage:")
        print("  python meeting_skill.py <audio_file>")
        print("  python meeting_skill.py --text <transcription_file>")
        print("  python meeting_skill.py --update <meeting_id> --title '新标题'")
        print("  python meeting_skill.py --query --date 2024-11-01")
        sys.exit(1)
    
    if sys.argv[1] == "--text":
        with open(sys.argv[2], "r", encoding="utf-8") as f:
            meeting = generate_minutes(f.read())
        print(json.dumps(meeting.to_dict(), ensure_ascii=False, indent=2))
        files = save_meeting(meeting)
        print(f"\nSaved to: {files}")
    
    elif sys.argv[1] == "--update":
        meeting_id = sys.argv[2]
        updates = {}
        for i in range(3, len(sys.argv), 2):
            if i + 1 < len(sys.argv) and sys.argv[i].startswith("--"):
                key = sys.argv[i][2:]
                updates[key] = sys.argv[i + 1]
        meeting = update_meeting(meeting_id, **updates)
        print(f"Updated to v{meeting.version}")
    
    elif sys.argv[1] == "--query":
        # 简单查询示例
        results = query_meetings()
        print(f"Found {len(results)} meetings")
        for r in results[:5]:
            print(f"  - {r['date']} {r['title']} (v{r['version']})")
    
    else:
        result = transcribe(sys.argv[1])
        meeting = generate_minutes(result["full_text"])
        print(json.dumps(meeting.to_dict(), ensure_ascii=False, indent=2))
        files = save_meeting(meeting)
        print(f"\nSaved to: {files}")
